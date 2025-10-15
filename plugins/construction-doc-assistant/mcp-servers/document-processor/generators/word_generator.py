"""
Word文档生成器

将Markdown文本转换为格式化的Word文档
"""
import os
import sys
from typing import Dict, Any, Optional, List
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    raise ImportError("请安装 python-docx: pip install python-docx")

from .base_generator import BaseGenerator
from .markdown_parser import MarkdownParser
from .construction_styles import ConstructionStyles
from utils import get_logger

logger = get_logger(__name__)


class WordGenerator(BaseGenerator):
    """Word文档生成器"""

    def __init__(self, template_type: str = "project_summary"):
        """
        初始化Word生成器

        Args:
            template_type: 模板类型
                - project_summary: 项目总结报告
                - inspection_report: 完整性检查报告
                - progress_analysis: 进度分析报告
                - organize_plan: 整理方案
        """
        super().__init__()
        self.template_type = template_type
        self.styles = ConstructionStyles.get_template(template_type)
        self.doc = None
        self.markdown_parser = MarkdownParser()
        self.warnings = []

    def generate(self,
                 markdown_file: str,
                 output_file: str,
                 options: Optional[Dict[str, Any]] = None) -> Dict:
        """
        生成Word文档

        Args:
            markdown_file: Markdown源文件路径
            output_file: Word输出文件路径
            options: 生成选项
                - project_info: 项目信息字典(用于页眉页脚)
                    - project_name: 项目名称
                    - report_type: 报告类型
                    - generate_date: 生成日期

        Returns:
            生成结果字典
        """
        try:
            # 1. 验证输入文件
            self.validate_input(markdown_file)
            self.validate_output(output_file)

            # 2. 读取Markdown文件
            self.logger.info(f"读取Markdown文件: {markdown_file}")
            with open(markdown_file, 'r', encoding='utf-8') as f:
                markdown_text = f.read()

            # 3. 解析Markdown
            self.logger.info("解析Markdown内容...")
            sections = self.markdown_parser.parse(markdown_text)
            self.logger.info(f"解析完成,共 {len(sections)} 个段落")

            # 4. 创建Word文档
            self.logger.info("创建Word文档...")
            self.doc = Document()
            self._setup_document_margins()

            # 5. 添加页眉页脚
            options = options or {}
            project_info = options.get('project_info')
            if project_info:
                self._add_header_footer(project_info)

            # 6. 逐节构建内容
            self.logger.info("构建Word内容...")
            for i, section in enumerate(sections):
                try:
                    self._add_section(section)
                except Exception as e:
                    self.logger.error(f"处理段落失败 (索引{i}): {e}", exc_info=True)
                    self.warnings.append(f"段落 {i+1} 处理失败: {section.get('type', 'unknown')} - {str(e)}")

            # 7. 保存文档
            self.logger.info(f"保存Word文档: {output_file}")
            self.doc.save(output_file)

            # 8. 返回结果
            result = self.create_success_response(
                output_file,
                sections_processed=len(sections),
                warnings=self.warnings,
                template_type=self.template_type,
            )

            self.logger.info(f"✅ Word文档生成成功: {output_file}")
            return result

        except Exception as e:
            self.logger.error(f"Word文档生成失败: {e}", exc_info=True)
            return self.create_error_response(str(e))

    def _setup_document_margins(self):
        """设置文档边距"""
        sections = self.doc.sections
        for section in sections:
            # 设置页边距(单位:Inches)
            section.top_margin = Inches(1.0)      # 上边距:1英寸
            section.bottom_margin = Inches(1.0)   # 下边距:1英寸
            section.left_margin = Inches(1.0)     # 左边距:1英寸
            section.right_margin = Inches(1.0)    # 右边距:1英寸

    def _add_header_footer(self, project_info: Dict):
        """
        添加页眉页脚

        Args:
            project_info: 项目信息
                - project_name: 项目名称
                - report_type: 报告类型
                - generate_date: 生成日期(可选)
        """
        section = self.doc.sections[0]

        # 页眉:项目名称 | 报告类型
        header = section.header
        header_para = header.paragraphs[0]
        header_text = f"{project_info.get('project_name', '')} - {project_info.get('report_type', '')}"
        header_para.text = header_text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 设置页眉样式
        for run in header_para.runs:
            run.font.size = Pt(10)
            run.font.name = ConstructionStyles.FONTS["body"]
            run.font.color.rgb = ConstructionStyles.COLORS["secondary"]

        # 页脚:页码
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码域
        self._add_page_number(footer_para)

    def _add_page_number(self, paragraph):
        """添加页码"""
        run = paragraph.add_run()
        run.font.size = Pt(10)
        run.font.name = ConstructionStyles.FONTS["body"]

        # 插入"第"
        run.add_text("第 ")

        # 插入页码域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        # 插入" 页"
        run.add_text(" 页")

    def _add_section(self, section: Dict):
        """
        添加一个段落

        Args:
            section: 段落数据字典
        """
        section_type = section.get("type")

        if section_type == "heading":
            self._add_heading(section)
        elif section_type == "paragraph":
            self._add_paragraph(section)
        elif section_type == "table":
            self._add_table(section)
        elif section_type == "list":
            self._add_list(section)
        elif section_type == "quote":
            self._add_quote(section)
        elif section_type == "code":
            self._add_code(section)
        elif section_type == "image":
            self._add_image_placeholder(section)
        elif section_type == "horizontal_rule":
            self._add_horizontal_rule()
        else:
            self.logger.warning(f"未知段落类型: {section_type}")

    def _add_heading(self, section: Dict):
        """添加标题"""
        level = section.get("level", 1)
        text = section.get("text", "")

        # 添加标题
        heading = self.doc.add_heading(text, level=min(level, 3))

        # 应用样式
        style_key = f"h{level}"
        if style_key in self.styles:
            self._apply_heading_style(heading, self.styles[style_key])

    def _add_paragraph(self, section: Dict):
        """添加段落"""
        text = section.get("text", "")

        # 添加段落
        para = self.doc.add_paragraph(text)

        # 应用样式
        if "body" in self.styles:
            self._apply_paragraph_style(para, self.styles["body"])

    def _add_table(self, section: Dict):
        """添加表格"""
        headers = section.get("headers", [])
        rows = section.get("rows", [])

        if not headers:
            self.logger.warning("表格缺少表头,跳过")
            return

        # 创建表格
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))

        # 应用表格样式
        table_style = self.styles.get("table", {}).get("style", "Light Grid Accent 1")
        table.style = table_style

        # 填充表头
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            if i < len(header_cells):
                cell = header_cells[i]
                cell.text = str(header_text)
                self._apply_table_header_style(cell)

        # 填充数据行
        for row_idx, row_data in enumerate(rows):
            if row_idx + 1 < len(table.rows):
                row_cells = table.rows[row_idx + 1].cells
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx < len(row_cells):
                        cell = row_cells[col_idx]
                        cell.text = str(cell_text)
                        self._apply_table_cell_style(cell)

    def _add_list(self, section: Dict):
        """添加列表"""
        ordered = section.get("ordered", False)
        items = section.get("items", [])

        list_style = self.styles.get("list", {})

        for i, item_text in enumerate(items):
            # 添加列表项
            if ordered:
                para = self.doc.add_paragraph(f"{i+1}. {item_text}")
            else:
                para = self.doc.add_paragraph(f"• {item_text}")

            # 应用样式
            self._apply_list_style(para, list_style)

    def _add_quote(self, section: Dict):
        """添加引用块"""
        text = section.get("text", "")

        # 添加引用段落
        para = self.doc.add_paragraph(text)

        # 应用引用样式
        if "quote" in self.styles:
            self._apply_quote_style(para, self.styles["quote"])

    def _add_code(self, section: Dict):
        """添加代码块"""
        code_text = section.get("text", "")
        language = section.get("language", "")

        # 如果有语言标识,添加为标签
        if language:
            label_para = self.doc.add_paragraph(f"[{language}]")
            label_run = label_para.runs[0]
            label_run.font.size = Pt(9)
            label_run.font.color.rgb = ConstructionStyles.COLORS["secondary"]
            label_para.paragraph_format.space_after = Pt(3)

        # 添加代码段落
        para = self.doc.add_paragraph(code_text)

        # 应用代码样式
        if "code" in self.styles:
            self._apply_code_style(para, self.styles["code"])

    def _add_image_placeholder(self, section: Dict):
        """
        添加图片占位符(Phase 1: 暂不实现图片)

        Args:
            section: 图片信息
        """
        alt_text = section.get("alt", "图片")
        image_url = section.get("url", "")

        # 添加占位符段落
        para = self.doc.add_paragraph()
        run = para.add_run(f"[图片: {alt_text}]")
        run.font.color.rgb = ConstructionStyles.COLORS["placeholder"]
        run.italic = True
        run.font.size = Pt(11)

        # 添加路径提示
        if image_url:
            hint_run = para.add_run(f"\n(路径: {image_url})")
            hint_run.font.size = Pt(9)
            hint_run.font.color.rgb = ConstructionStyles.COLORS["placeholder"]

        para.paragraph_format.space_after = Pt(12)

        # 记录警告
        self.warnings.append(f"图片暂未支持,已添加占位符: {alt_text}")
        self.logger.info(f"📷 图片占位符已添加: {alt_text} (Phase 2将实现)")

        # TODO Phase 2: 实现图片插入
        """
        from PIL import Image

        # 1. 解析图片路径
        image_path = self._resolve_image_path(image_url)

        # 2. 验证图片存在
        if not os.path.exists(image_path):
            self._add_image_error_placeholder(alt_text, image_url)
            return

        # 3. 插入图片
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(5))

        # 4. 添加题注
        caption = self.doc.add_paragraph(alt_text)
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        """

    def _add_horizontal_rule(self):
        """添加水平线"""
        para = self.doc.add_paragraph()
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)

        # 使用连续的下划线模拟水平线
        run = para.add_run("_" * 60)
        run.font.color.rgb = ConstructionStyles.COLORS["secondary"]
        run.font.size = Pt(8)

    # ========== 样式应用方法 ==========

    def _apply_heading_style(self, paragraph, style_config: Dict):
        """应用标题样式"""
        # 字体和大小
        for run in paragraph.runs:
            if "font_name" in style_config:
                run.font.name = style_config["font_name"]
            if "font_size" in style_config:
                run.font.size = Pt(style_config["font_size"])
            if "bold" in style_config:
                run.font.bold = style_config["bold"]
            if "color" in style_config:
                color_name = style_config["color"]
                run.font.color.rgb = ConstructionStyles.get_color(color_name)

        # 对齐方式
        if "alignment" in style_config:
            alignment = style_config["alignment"]
            if alignment in ConstructionStyles.ALIGNMENT_MAP:
                paragraph.alignment = ConstructionStyles.ALIGNMENT_MAP[alignment]

        # 段前段后间距
        if "space_before" in style_config:
            paragraph.paragraph_format.space_before = Pt(style_config["space_before"])
        if "space_after" in style_config:
            paragraph.paragraph_format.space_after = Pt(style_config["space_after"])

    def _apply_paragraph_style(self, paragraph, style_config: Dict):
        """应用段落样式"""
        # 字体和大小
        for run in paragraph.runs:
            if "font_name" in style_config:
                run.font.name = style_config["font_name"]
            if "font_size" in style_config:
                run.font.size = Pt(style_config["font_size"])
            if "color" in style_config:
                color_name = style_config["color"]
                run.font.color.rgb = ConstructionStyles.get_color(color_name)

        # 行距
        if "line_spacing" in style_config:
            paragraph.paragraph_format.line_spacing = style_config["line_spacing"]

        # 对齐方式
        if "alignment" in style_config:
            alignment = style_config["alignment"]
            if alignment in ConstructionStyles.ALIGNMENT_MAP:
                paragraph.alignment = ConstructionStyles.ALIGNMENT_MAP[alignment]

        # 首行缩进
        if "first_line_indent" in style_config:
            indent_value = style_config["first_line_indent"]
            if isinstance(indent_value, (int, float)) and indent_value > 0:
                paragraph.paragraph_format.first_line_indent = Inches(indent_value)

        # 段后间距
        if "space_after" in style_config:
            paragraph.paragraph_format.space_after = Pt(style_config["space_after"])

    def _apply_table_header_style(self, cell):
        """应用表头样式"""
        table_config = self.styles.get("table", {})

        # 设置背景色
        if "header_bg" in table_config:
            bg_color = table_config["header_bg"]
            self._set_cell_background(cell, bg_color)

        # 设置字体
        for para in cell.paragraphs:
            for run in para.runs:
                if "header_font_name" in table_config:
                    run.font.name = table_config["header_font_name"]
                if "header_font_size" in table_config:
                    run.font.size = Pt(table_config["header_font_size"])
                if "header_bold" in table_config:
                    run.font.bold = table_config["header_bold"]

            # 对齐方式
            if "header_alignment" in table_config:
                alignment = table_config["header_alignment"]
                if alignment in ConstructionStyles.ALIGNMENT_MAP:
                    para.alignment = ConstructionStyles.ALIGNMENT_MAP[alignment]

    def _apply_table_cell_style(self, cell):
        """应用表格单元格样式"""
        table_config = self.styles.get("table", {})

        # 设置字体
        for para in cell.paragraphs:
            for run in para.runs:
                if "cell_font_name" in table_config:
                    run.font.name = table_config["cell_font_name"]
                if "cell_font_size" in table_config:
                    run.font.size = Pt(table_config["cell_font_size"])

            # 对齐方式
            if "cell_alignment" in table_config:
                alignment = table_config["cell_alignment"]
                if alignment in ConstructionStyles.ALIGNMENT_MAP:
                    para.alignment = ConstructionStyles.ALIGNMENT_MAP[alignment]

    def _apply_list_style(self, paragraph, style_config: Dict):
        """应用列表样式"""
        for run in paragraph.runs:
            if "font_name" in style_config:
                run.font.name = style_config["font_name"]
            if "font_size" in style_config:
                run.font.size = Pt(style_config["font_size"])

        if "line_spacing" in style_config:
            paragraph.paragraph_format.line_spacing = style_config["line_spacing"]

    def _apply_quote_style(self, paragraph, style_config: Dict):
        """应用引用样式"""
        for run in paragraph.runs:
            if "font_name" in style_config:
                run.font.name = style_config["font_name"]
            if "font_size" in style_config:
                run.font.size = Pt(style_config["font_size"])
            if "italic" in style_config:
                run.italic = style_config["italic"]
            if "color" in style_config:
                color_name = style_config["color"]
                run.font.color.rgb = ConstructionStyles.get_color(color_name)

        if "left_indent" in style_config:
            paragraph.paragraph_format.left_indent = style_config["left_indent"]

        if "space_after" in style_config:
            paragraph.paragraph_format.space_after = Pt(style_config["space_after"])

    def _apply_code_style(self, paragraph, style_config: Dict):
        """应用代码样式"""
        for run in paragraph.runs:
            if "font_name" in style_config:
                run.font.name = style_config["font_name"]
            if "font_size" in style_config:
                run.font.size = Pt(style_config["font_size"])

        if "line_spacing" in style_config:
            paragraph.paragraph_format.line_spacing = style_config["line_spacing"]

        if "left_indent" in style_config:
            paragraph.paragraph_format.left_indent = style_config["left_indent"]

    def _set_cell_background(self, cell, color_hex: str):
        """
        设置单元格背景色

        Args:
            cell: 单元格对象
            color_hex: 十六进制颜色(如"D9E2F3")
        """
        cell_xml = cell._element
        cell_properties = cell_xml.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), color_hex)
        cell_properties.append(shade_obj)
