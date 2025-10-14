"""
Word 文档解析器模块

解析 .docx 格式的 Word 文档
提取文本、表格、标题结构等信息
"""
from typing import Dict, List, Optional, Any
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from .base_parser import BaseParser
from utils import get_logger, ParseError, config

logger = get_logger(__name__)


class WordParser(BaseParser):
    """Word 文档解析器"""

    def __init__(self):
        super().__init__()

    def get_supported_extensions(self) -> list:
        """获取支持的文件扩展名"""
        return ['.docx', '.doc']

    def parse(self, file_path: str, options: Optional[Dict[str, Any]] = None) -> Dict:
        """
        解析 Word 文档

        Args:
            file_path: Word 文档路径
            options: 解析选项
                - extract_tables: 是否提取表格 (默认 True)
                - max_paragraphs: 最大段落数限制
                - keywords: 关注的关键词列表

        Returns:
            解析结果字典
        """
        try:
            from docx import Document
        except ImportError:
            raise ParseError(
                "缺少 python-docx 库，请运行: pip install python-docx"
            )

        options = options or {}
        extract_tables = options.get('extract_tables', True)
        max_paragraphs = options.get('max_paragraphs', None)
        keywords = options.get('keywords', [])

        try:
            # 加载文档
            self.logger.info(f"加载 Word 文档: {file_path}")
            doc = Document(file_path)

            # 提取内容
            content = {}

            # 1. 提取章节和段落
            sections = self._extract_sections(doc, max_paragraphs, keywords)
            content['sections'] = sections

            # 2. 提取表格
            if extract_tables:
                tables = self._extract_tables(doc)
                content['tables'] = tables
            else:
                content['tables'] = []

            # 3. 提取文档大纲
            outline = self._extract_outline(doc)
            content['outline'] = outline

            # 4. 生成摘要
            summary = self._generate_summary(sections, content['tables'])

            # 5. 提取元数据
            metadata = self._extract_metadata(doc)

            return self._create_success_response(
                file_path,
                content,
                summary,
                metadata
            )

        except Exception as e:
            self.logger.error(f"Word 文档解析失败: {e}", exc_info=True)
            raise ParseError(f"Word 文档解析失败: {str(e)}")

    def _extract_sections(
        self,
        doc,
        max_paragraphs: Optional[int] = None,
        keywords: Optional[List[str]] = None
    ) -> Dict[str, List[str]]:
        """
        提取章节和段落，按标题分组

        Args:
            doc: Document 对象
            max_paragraphs: 最大段落数
            keywords: 关键词列表

        Returns:
            章节字典 {章节名: [段落列表]}
        """
        sections = {}
        current_section = "文档开头"  # 默认章节名
        sections[current_section] = []

        paragraph_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()

            # 跳过空段落
            if not text:
                continue

            # 检查是否为标题
            if para.style.name.startswith('Heading'):
                # 创建新章节
                current_section = text
                sections[current_section] = []
                self.logger.debug(f"发现章节: {current_section}")
            else:
                # 添加到当前章节
                # 如果指定了关键词，只保留包含关键词的段落
                if keywords:
                    if any(kw in text for kw in keywords):
                        sections[current_section].append(text)
                        paragraph_count += 1
                else:
                    sections[current_section].append(text)
                    paragraph_count += 1

            # 检查是否达到最大段落数
            if max_paragraphs and paragraph_count >= max_paragraphs:
                self.logger.info(f"已达到最大段落数限制: {max_paragraphs}")
                break

        self.logger.info(f"提取章节: {len(sections)} 个, 段落: {paragraph_count} 个")
        return sections

    def _extract_tables(self, doc) -> List[Dict]:
        """
        提取表格数据

        Args:
            doc: Document 对象

        Returns:
            表格列表
        """
        tables = []

        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns),
                "data": [],
                "headers": []
            }

            # 提取表格内容
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)

                table_data["data"].append(row_data)

            # 智能识别表头（第一行通常是表头）
            if table_data["data"]:
                table_data["headers"] = table_data["data"][0]

            tables.append(table_data)

        self.logger.info(f"提取表格: {len(tables)} 个")
        return tables

    def _extract_outline(self, doc) -> List[Dict]:
        """
        提取文档大纲（标题结构）

        Args:
            doc: Document 对象

        Returns:
            大纲列表
        """
        outline = []

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = self._get_heading_level(para.style.name)
                outline.append({
                    "level": level,
                    "text": para.text.strip()
                })

        self.logger.info(f"提取大纲: {len(outline)} 个标题")
        return outline

    @staticmethod
    def _get_heading_level(style_name: str) -> int:
        """
        从样式名称提取标题级别

        Args:
            style_name: 样式名称，如 'Heading 1'

        Returns:
            标题级别 (1-9)
        """
        try:
            # 提取数字部分
            parts = style_name.split()
            if len(parts) >= 2 and parts[-1].isdigit():
                return int(parts[-1])
        except:
            pass
        return 1  # 默认为一级标题

    def _extract_metadata(self, doc) -> Dict:
        """
        提取文档元数据

        Args:
            doc: Document 对象

        Returns:
            元数据字典
        """
        metadata = {}

        try:
            core_props = doc.core_properties

            if hasattr(core_props, 'author') and core_props.author:
                metadata['author'] = core_props.author

            if hasattr(core_props, 'title') and core_props.title:
                metadata['title'] = core_props.title

            if hasattr(core_props, 'subject') and core_props.subject:
                metadata['subject'] = core_props.subject

            if hasattr(core_props, 'created') and core_props.created:
                metadata['created'] = str(core_props.created)

            if hasattr(core_props, 'modified') and core_props.modified:
                metadata['modified'] = str(core_props.modified)

        except Exception as e:
            self.logger.warning(f"提取元数据失败: {e}")

        return metadata

    def _generate_summary(self, sections: Dict, tables: List) -> Dict:
        """
        生成摘要信息

        Args:
            sections: 章节字典
            tables: 表格列表

        Returns:
            摘要字典
        """
        # 统计段落数
        total_paragraphs = sum(len(paragraphs) for paragraphs in sections.values())

        # 统计总字数（所有段落的字符数）
        total_chars = sum(
            sum(len(p) for p in paragraphs)
            for paragraphs in sections.values()
        )

        # 提取章节列表
        section_titles = [
            title for title in sections.keys()
            if title != "文档开头"
        ]

        summary = {
            "total_sections": len(sections),
            "total_paragraphs": total_paragraphs,
            "total_chars": total_chars,
            "total_tables": len(tables),
            "section_titles": section_titles
        }

        # 如果有表格，添加表格统计
        if tables:
            total_table_rows = sum(t["rows"] for t in tables)
            summary["total_table_rows"] = total_table_rows

        return summary

    def extract_text_by_keywords(
        self,
        file_path: str,
        keywords: List[str]
    ) -> Dict:
        """
        根据关键词提取相关文本

        Args:
            file_path: 文档路径
            keywords: 关键词列表

        Returns:
            包含关键词的文本片段
        """
        result = self.parse(file_path, {'keywords': keywords})

        if result['status'] != 'success':
            return result

        # 过滤空章节
        filtered_sections = {
            section: paragraphs
            for section, paragraphs in result['content']['sections'].items()
            if paragraphs
        }

        result['content']['sections'] = filtered_sections
        result['content']['keyword_filter'] = keywords

        return result
