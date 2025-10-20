#!/usr/bin/env python3
"""
NotebookLM Report Generator MCP Server
专业报告生成服务 - Word/PDF 排版和导出
"""

import os
import json
import logging
import sys
from pathlib import Path
from datetime import datetime
from typing import Any, List, Dict

# MCP imports
from mcp.server import Server
from mcp.types import Tool, TextContent

# 文档处理库
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    stream=sys.stderr,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 创建 MCP 服务器实例
server = Server("notebooklm-report-generator")

@server.list_tools()
async def list_tools() -> List[Tool]:
    """列出所有可用工具"""
    return [
        Tool(
            name="generate_word_report",
            description="生成专业排版的 Word 报告（支持多级标题、表格、图表占位符）",
            inputSchema={
                "type": "object",
                "properties": {
                    "template": {
                        "type": "string",
                        "description": "模板名称（如 'research-report.docx'）或 'default'",
                        "default": "default"
                    },
                    "content": {
                        "type": "object",
                        "description": "报告内容（结构化数据）",
                        "properties": {
                            "title": {"type": "string"},
                            "author": {"type": "string"},
                            "date": {"type": "string"},
                            "abstract": {"type": "string"},
                            "chapters": {"type": "array"},
                            "tables": {"type": "array"},
                            "charts": {"type": "array"},
                            "references": {"type": "array"}
                        }
                    },
                    "style": {
                        "type": "string",
                        "enum": ["academic", "business", "technical"],
                        "description": "报告风格",
                        "default": "business"
                    },
                    "output_path": {
                        "type": "string",
                        "description": "输出文件路径（绝对路径）"
                    }
                },
                "required": ["content", "output_path"]
            }
        ),
        Tool(
            name="insert_table",
            description="在 Word 文档中插入专业表格",
            inputSchema={
                "type": "object",
                "properties": {
                    "docx_path": {
                        "type": "string",
                        "description": "Word 文档路径"
                    },
                    "table_data": {
                        "type": "object",
                        "properties": {
                            "title": {"type": "string"},
                            "headers": {"type": "array"},
                            "rows": {"type": "array"}
                        }
                    },
                    "position": {
                        "type": "string",
                        "description": "插入位置（章节号，如 '2.1'）"
                    }
                },
                "required": ["docx_path", "table_data"]
            }
        ),
        Tool(
            name="list_templates",
            description="列出所有可用的报告模板",
            inputSchema={
                "type": "object",
                "properties": {}
            }
        ),
        Tool(
            name="convert_to_pdf",
            description="将 Word 文档转换为 PDF（需要 LibreOffice）",
            inputSchema={
                "type": "object",
                "properties": {
                    "docx_path": {
                        "type": "string",
                        "description": "Word 文档路径"
                    },
                    "pdf_path": {
                        "type": "string",
                        "description": "PDF 输出路径（可选）"
                    }
                },
                "required": ["docx_path"]
            }
        )
    ]

@server.call_tool()
async def call_tool(name: str, arguments: Any) -> List[TextContent]:
    """工具调用路由"""
    try:
        if not DOCX_AVAILABLE:
            return [TextContent(
                type="text",
                text=json.dumps({
                    "error": "python-docx 库未安装，请运行: pip install python-docx"
                }, ensure_ascii=False)
            )]

        if name == "generate_word_report":
            return await generate_word_report(
                arguments.get("template", "default"),
                arguments["content"],
                arguments.get("style", "business"),
                arguments["output_path"]
            )
        elif name == "insert_table":
            return await insert_table(
                arguments["docx_path"],
                arguments["table_data"],
                arguments.get("position")
            )
        elif name == "list_templates":
            return await list_templates()
        elif name == "convert_to_pdf":
            return await convert_to_pdf(
                arguments["docx_path"],
                arguments.get("pdf_path")
            )
        else:
            raise ValueError(f"未知工具: {name}")

    except Exception as e:
        logger.error(f"工具执行错误: {e}", exc_info=True)
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": str(e),
                "type": type(e).__name__
            }, ensure_ascii=False)
        )]

async def generate_word_report(
    template: str,
    content: Dict[str, Any],
    style: str,
    output_path: str
) -> List[TextContent]:
    """生成 Word 报告"""
    logger.info(f"生成报告: 模板={template}, 风格={style}, 输出={output_path}")

    # 1. 加载模板或创建新文档
    if template == "default":
        doc = Document()
        apply_default_styles(doc, style)
    else:
        template_dir = Path(__file__).parent / "templates"
        template_path = template_dir / template
        if template_path.exists():
            doc = Document(str(template_path))
        else:
            logger.warning(f"模板不存在: {template}，使用默认模板")
            doc = Document()
            apply_default_styles(doc, style)

    # 2. 添加封面
    add_cover_page(
        doc,
        content.get("title", "未命名报告"),
        content.get("author", "NotebookLM Assistant"),
        content.get("date", datetime.now().strftime("%Y-%m-%d"))
    )

    # 3. 添加摘要（如果有）
    if "abstract" in content and content["abstract"]:
        doc.add_page_break()
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(content["abstract"])

    # 4. 添加目录占位符
    doc.add_page_break()
    heading = doc.add_heading("目录", level=1)
    p = doc.add_paragraph("[此处应插入自动目录，请在 Word 中使用\"引用\"-\"目录\"功能更新]")
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # 5. 逐章节添加内容
    for chapter in content.get("chapters", []):
        doc.add_page_break()

        # 章节标题
        doc.add_heading(chapter.get("title", "未命名章节"), level=1)

        # 章节下的小节
        for section in chapter.get("sections", []):
            # 小节标题
            doc.add_heading(section.get("title", "未命名小节"), level=2)

            # 段落内容
            section_content = section.get("content", "")
            if section_content:
                paragraphs = section_content.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())

            # 插入表格（如果有）
            if "table" in section:
                add_table_to_doc(doc, section["table"])

            # 插入图表占位符（如果有）
            if "chart" in section:
                add_chart_placeholder(doc, section["chart"])

    # 6. 添加参考文献
    if content.get("references"):
        doc.add_page_break()
        doc.add_heading("参考文献", level=1)
        for i, ref in enumerate(content["references"], 1):
            citation = ref.get("formatted_citation", str(ref))
            doc.add_paragraph(f"[{i}] {citation}")

    # 7. 保存文档
    output_path_obj = Path(output_path)
    output_path_obj.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path_obj))

    logger.info(f"报告生成完成: {output_path}")

    return [TextContent(
        type="text",
        text=json.dumps({
            "status": "success",
            "output_path": str(output_path_obj),
            "pages": len([s for s in doc.sections]),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables)
        }, ensure_ascii=False, indent=2)
    )]

def apply_default_styles(doc: Document, style: str):
    """应用默认样式"""
    styles = doc.styles

    # 设置正文样式
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # 中文字体
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 标题样式
    if style == "academic":
        # 学术风格：黑色、加粗、居中
        for level in range(1, 4):
            heading_style = styles[f'Heading {level}']
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.font.bold = True

    elif style == "business":
        # 商业风格：深蓝色、加粗
        for level in range(1, 4):
            heading_style = styles[f'Heading {level}']
            heading_style.font.color.rgb = RGBColor(0, 70, 127)  # 深蓝色
            heading_style.font.bold = True

    elif style == "technical":
        # 技术风格：深灰色、加粗
        for level in range(1, 4):
            heading_style = styles[f'Heading {level}']
            heading_style.font.color.rgb = RGBColor(64, 64, 64)  # 深灰色
            heading_style.font.bold = True

def add_cover_page(doc: Document, title: str, author: str, date: str):
    """添加封面"""
    # 标题
    p_title = doc.add_paragraph()
    run = p_title.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 空行
    for _ in range(10):
        doc.add_paragraph()

    # 作者和日期
    p_author = doc.add_paragraph()
    p_author.add_run(f"生成者: {author}\n")
    p_author.add_run(f"日期: {date}")
    p_author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_table_to_doc(doc: Document, table_data: Dict[str, Any]):
    """添加表格到文档"""
    headers = table_data.get("headers", [])
    rows = table_data.get("rows", [])

    if not headers or not rows:
        return

    # 添加表格标题
    if "title" in table_data and table_data["title"]:
        p = doc.add_paragraph()
        run = p.add_run(table_data["title"])
        run.font.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 创建表格
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # 填充表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = str(header)
        # 表头加粗
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # 填充数据行
    for i, row_data in enumerate(rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)

    doc.add_paragraph()  # 表格后空行

def add_chart_placeholder(doc: Document, chart_data: Dict[str, Any]):
    """添加图表占位符"""
    p = doc.add_paragraph()
    run = p.add_run(f"[图表: {chart_data.get('title', '未命名图表')}]")
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # 添加说明
    if "description" in chart_data and chart_data["description"]:
        doc.add_paragraph(chart_data["description"])

    # 添加数据表（供用户手动创建图表）
    if "data" in chart_data:
        data = chart_data["data"]
        if "labels" in data and "datasets" in data:
            table_data = {
                "title": f"{chart_data['title']} - 数据",
                "headers": data["labels"],
                "rows": [dataset.get("data", []) for dataset in data["datasets"]]
            }
            add_table_to_doc(doc, table_data)

async def insert_table(
    docx_path: str,
    table_data: Dict[str, Any],
    position: str = None
) -> List[TextContent]:
    """在现有文档中插入表格"""
    logger.info(f"插入表格到: {docx_path}, 位置: {position}")

    if not Path(docx_path).exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"文档不存在: {docx_path}"
            }, ensure_ascii=False)
        )]

    doc = Document(docx_path)
    add_table_to_doc(doc, table_data)
    doc.save(docx_path)

    return [TextContent(
        type="text",
        text=json.dumps({
            "status": "success",
            "message": "表格已插入"
        }, ensure_ascii=False)
    )]

async def list_templates() -> List[TextContent]:
    """列出所有可用模板"""
    template_dir = Path(__file__).parent / "templates"
    template_dir.mkdir(exist_ok=True)

    templates = list(template_dir.glob("*.docx"))

    return [TextContent(
        type="text",
        text=json.dumps({
            "templates": [t.name for t in templates],
            "default_available": True,
            "template_directory": str(template_dir)
        }, ensure_ascii=False, indent=2)
    )]

async def convert_to_pdf(
    docx_path: str,
    pdf_path: str = None
) -> List[TextContent]:
    """转换 Word 文档为 PDF"""
    logger.info(f"转换为 PDF: {docx_path}")

    if not Path(docx_path).exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"文档不存在: {docx_path}"
            }, ensure_ascii=False)
        )]

    if pdf_path is None:
        pdf_path = str(Path(docx_path).with_suffix('.pdf'))

    try:
        import subprocess

        # 使用 LibreOffice 转换
        result = subprocess.run([
            'soffice',
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            str(Path(pdf_path).parent),
            docx_path
        ], capture_output=True, text=True, timeout=60)

        if result.returncode == 0:
            return [TextContent(
                type="text",
                text=json.dumps({
                    "status": "success",
                    "pdf_path": str(pdf_path)
                }, ensure_ascii=False)
            )]
        else:
            return [TextContent(
                type="text",
                text=json.dumps({
                    "error": "PDF 转换失败",
                    "details": result.stderr
                }, ensure_ascii=False)
            )]

    except FileNotFoundError:
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": "未找到 LibreOffice，请安装: brew install libreoffice (macOS)"
            }, ensure_ascii=False)
        )]
    except Exception as e:
        logger.error(f"PDF 转换错误: {e}", exc_info=True)
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": str(e)
            }, ensure_ascii=False)
        )]

async def main():
    """启动 MCP 服务器"""
    from mcp.server.stdio import stdio_server

    logger.info("启动 NotebookLM Report Generator MCP Server...")

    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
