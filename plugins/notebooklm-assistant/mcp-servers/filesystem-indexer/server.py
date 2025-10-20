#!/usr/bin/env python3
"""
NotebookLM Filesystem Indexer MCP Server
提供轻量级文档索引和按需读取功能
"""

import os
import json
import logging
import sys
from pathlib import Path
from typing import Any, List, Dict
from datetime import datetime

# MCP imports
from mcp.server import Server
from mcp.types import Tool, TextContent

# 配置日志 - 只写到 stderr
logging.basicConfig(
    level=logging.INFO,
    stream=sys.stderr,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 创建 MCP 服务器实例
server = Server("notebooklm-filesystem")

@server.list_tools()
async def list_tools() -> List[Tool]:
    """列出所有可用工具"""
    return [
        Tool(
            name="scan_directory",
            description="扫描目录生成轻量级索引（仅元数据，不读取文件内容）",
            inputSchema={
                "type": "object",
                "properties": {
                    "directory": {
                        "type": "string",
                        "description": "要扫描的目录路径（绝对路径）"
                    },
                    "file_types": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "文件类型过滤（如 ['.pdf', '.docx']），留空表示全部",
                        "default": []
                    },
                    "max_depth": {
                        "type": "integer",
                        "description": "最大扫描深度（默认 10）",
                        "default": 10
                    }
                },
                "required": ["directory"]
            }
        ),
        Tool(
            name="preview_document",
            description="预览文档前 N 个字符（用于相关性评估）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档的绝对路径"
                    },
                    "preview_length": {
                        "type": "integer",
                        "description": "预览长度（字符数，默认 500）",
                        "default": 500
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="parse_document_smart",
            description="智能解析文档，支持针对性提取（基于关键词）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档的绝对路径"
                    },
                    "mode": {
                        "type": "string",
                        "enum": ["full", "targeted", "summary"],
                        "description": "full=完整内容, targeted=关键词定向提取, summary=仅摘要"
                    },
                    "keywords": {
                        "type": "array",
                        "items": {"type": "string"},
                        "description": "关键词列表（targeted 模式必需）",
                        "default": []
                    }
                },
                "required": ["file_path", "mode"]
            }
        ),
        Tool(
            name="extract_keywords",
            description="从文档中提取关键词和主题（用于索引增强）",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档的绝对路径"
                    },
                    "top_k": {
                        "type": "integer",
                        "description": "返回前 K 个关键词（默认 20）",
                        "default": 20
                    }
                },
                "required": ["file_path"]
            }
        )
    ]

@server.call_tool()
async def call_tool(name: str, arguments: Any) -> List[TextContent]:
    """工具调用路由"""
    try:
        if name == "scan_directory":
            return await scan_directory(
                arguments["directory"],
                arguments.get("file_types", []),
                arguments.get("max_depth", 10)
            )
        elif name == "preview_document":
            return await preview_document(
                arguments["file_path"],
                arguments.get("preview_length", 500)
            )
        elif name == "parse_document_smart":
            return await parse_document_smart(
                arguments["file_path"],
                arguments["mode"],
                arguments.get("keywords", [])
            )
        elif name == "extract_keywords":
            return await extract_keywords(
                arguments["file_path"],
                arguments.get("top_k", 20)
            )
        else:
            raise ValueError(f"未知工具: {name}")
    except Exception as e:
        logger.error(f"工具执行错误: {e}", exc_info=True)
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": str(e),
                "type": type(e).__name__
            }, ensure_ascii=False)
        )]

async def scan_directory(
    directory: str,
    file_types: List[str],
    max_depth: int
) -> List[TextContent]:
    """扫描目录生成轻量级索引"""
    logger.info(f"扫描目录: {directory}, 文件类型: {file_types}, 最大深度: {max_depth}")

    index = []
    base_path = Path(directory)

    if not base_path.exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"目录不存在: {directory}"
            }, ensure_ascii=False)
        )]

    # 支持的文档类型
    supported_types = {
        '.pdf', '.doc', '.docx', '.xls', '.xlsx',
        '.ppt', '.pptx', '.txt', '.md', '.csv'
    }

    for root, dirs, files in os.walk(base_path):
        # 计算当前深度
        try:
            depth = len(Path(root).relative_to(base_path).parts)
        except ValueError:
            depth = 0

        if depth > max_depth:
            continue

        for file in files:
            file_path = Path(root) / file
            ext = file_path.suffix.lower()

            # 文件类型过滤
            if file_types and ext not in file_types:
                continue

            # 只索引支持的文档类型
            if ext not in supported_types:
                continue

            # 收集元数据（不读取文件内容）
            try:
                stat = file_path.stat()
                index.append({
                    "path": str(file_path),
                    "name": file_path.name,
                    "extension": ext,
                    "size": stat.st_size,
                    "modified": stat.st_mtime,
                    "modified_readable": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                    "relative_path": str(file_path.relative_to(base_path))
                })
            except Exception as e:
                logger.warning(f"无法读取文件信息: {file_path}, 错误: {e}")
                continue

    # 按修改时间倒序排列（最新的在前）
    index.sort(key=lambda x: x["modified"], reverse=True)

    result = {
        "total_files": len(index),
        "file_types": {},
        "index": index
    }

    # 统计文件类型分布
    for item in index:
        ext = item["extension"]
        result["file_types"][ext] = result["file_types"].get(ext, 0) + 1

    logger.info(f"扫描完成，找到 {len(index)} 个文档")

    return [TextContent(
        type="text",
        text=json.dumps(result, ensure_ascii=False, indent=2)
    )]

async def preview_document(
    file_path: str,
    preview_length: int
) -> List[TextContent]:
    """预览文档前 N 个字符"""
    logger.info(f"预览文档: {file_path}, 长度: {preview_length}")

    path = Path(file_path)
    if not path.exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"文件不存在: {file_path}"
            }, ensure_ascii=False)
        )]

    ext = path.suffix.lower()
    preview = ""

    try:
        if ext == '.txt' or ext == '.md':
            # 纯文本文件
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                preview = f.read(preview_length)

        elif ext == '.pdf':
            # PDF 文件 - 提取第一页文本
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    if len(reader.pages) > 0:
                        preview = reader.pages[0].extract_text()[:preview_length]
                    else:
                        preview = "[PDF 文件无内容]"
            except ImportError:
                preview = "[需要安装 PyPDF2 库来预览 PDF]"
            except Exception as e:
                preview = f"[PDF 读取错误: {str(e)}]"

        elif ext in ['.docx', '.doc']:
            # Word 文件
            try:
                from docx import Document
                doc = Document(file_path)
                text = '\n'.join([p.text for p in doc.paragraphs[:3]])
                preview = text[:preview_length]
            except ImportError:
                preview = "[需要安装 python-docx 库来预览 Word 文档]"
            except Exception as e:
                preview = f"[Word 读取错误: {str(e)}]"

        elif ext in ['.xlsx', '.xls']:
            # Excel 文件
            try:
                import openpyxl
                wb = openpyxl.load_workbook(file_path, read_only=True)
                ws = wb.active
                rows = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 5:  # 只读前5行
                        break
                    rows.append(str(row))
                preview = '\n'.join(rows)[:preview_length]
            except ImportError:
                preview = "[需要安装 openpyxl 库来预览 Excel 文件]"
            except Exception as e:
                preview = f"[Excel 读取错误: {str(e)}]"

        else:
            preview = f"[不支持预览的文件类型: {ext}]"

    except Exception as e:
        logger.error(f"预览文档错误: {e}", exc_info=True)
        preview = f"[预览错误: {str(e)}]"

    return [TextContent(
        type="text",
        text=json.dumps({
            "file_path": file_path,
            "preview_length": len(preview),
            "preview": preview
        }, ensure_ascii=False, indent=2)
    )]

async def parse_document_smart(
    file_path: str,
    mode: str,
    keywords: List[str]
) -> List[TextContent]:
    """智能文档解析"""
    logger.info(f"智能解析: {file_path}, 模式: {mode}, 关键词: {keywords}")

    path = Path(file_path)
    if not path.exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"文件不存在: {file_path}"
            }, ensure_ascii=False)
        )]

    ext = path.suffix.lower()
    result = {
        "file_path": file_path,
        "mode": mode,
        "keywords": keywords,
        "content": "",
        "matched_sections": []
    }

    try:
        if mode == "full":
            # 完整内容模式
            if ext == '.txt' or ext == '.md':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    result["content"] = f.read()
            elif ext == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = []
                    for page in reader.pages:
                        text.append(page.extract_text())
                    result["content"] = '\n\n'.join(text)
            elif ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(file_path)
                result["content"] = '\n'.join([p.text for p in doc.paragraphs])
            else:
                result["content"] = f"[不支持的文件类型: {ext}]"

        elif mode == "targeted":
            # 针对性提取模式（只提取包含关键词的段落）
            if not keywords:
                return [TextContent(
                    type="text",
                    text=json.dumps({
                        "error": "targeted 模式需要提供关键词"
                    }, ensure_ascii=False)
                )]

            # 读取全文
            full_text = ""
            if ext == '.txt' or ext == '.md':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    full_text = f.read()
            elif ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(file_path)
                full_text = '\n'.join([p.text for p in doc.paragraphs])

            # 分段并查找包含关键词的段落
            paragraphs = full_text.split('\n')
            for i, para in enumerate(paragraphs):
                for keyword in keywords:
                    if keyword.lower() in para.lower():
                        # 提取当前段落 + 上下文（前后各1段）
                        context_start = max(0, i - 1)
                        context_end = min(len(paragraphs), i + 2)
                        context = '\n'.join(paragraphs[context_start:context_end])

                        result["matched_sections"].append({
                            "keyword": keyword,
                            "paragraph_index": i,
                            "content": context
                        })
                        break  # 每个段落只匹配一次

        elif mode == "summary":
            # 摘要模式（只返回开头部分）
            if ext == '.txt' or ext == '.md':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    result["content"] = f.read(2000)  # 前2000字符
            elif ext in ['.docx', '.doc']:
                from docx import Document
                doc = Document(file_path)
                paragraphs = [p.text for p in doc.paragraphs[:5]]  # 前5段
                result["content"] = '\n'.join(paragraphs)
            else:
                result["content"] = f"[不支持的文件类型: {ext}]"

    except Exception as e:
        logger.error(f"解析文档错误: {e}", exc_info=True)
        result["error"] = str(e)

    return [TextContent(
        type="text",
        text=json.dumps(result, ensure_ascii=False, indent=2)
    )]

async def extract_keywords(
    file_path: str,
    top_k: int
) -> List[TextContent]:
    """提取关键词（简单实现：基于词频）"""
    logger.info(f"提取关键词: {file_path}, top_k: {top_k}")

    path = Path(file_path)
    if not path.exists():
        return [TextContent(
            type="text",
            text=json.dumps({
                "error": f"文件不存在: {file_path}"
            }, ensure_ascii=False)
        )]

    # 简单实现：读取文本并统计词频
    # TODO: 可以集成 jieba（中文）或 NLTK（英文）进行更好的分词

    ext = path.suffix.lower()
    text = ""

    try:
        if ext == '.txt' or ext == '.md':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        elif ext in ['.docx', '.doc']:
            from docx import Document
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])

        # 简单的词频统计（按空格分词）
        words = text.lower().split()
        word_freq = {}

        # 停用词列表（简化版）
        stopwords = {'的', '了', '是', '在', '和', '有', '与', '等', '中', '及',
                     'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to'}

        for word in words:
            # 过滤短词和停用词
            if len(word) > 2 and word not in stopwords:
                word_freq[word] = word_freq.get(word, 0) + 1

        # 按频率排序
        sorted_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [word for word, freq in sorted_keywords[:top_k]]

    except Exception as e:
        logger.error(f"提取关键词错误: {e}", exc_info=True)
        keywords = []

    return [TextContent(
        type="text",
        text=json.dumps({
            "file_path": file_path,
            "keywords": keywords
        }, ensure_ascii=False, indent=2)
    )]

async def main():
    """启动 MCP 服务器"""
    from mcp.server.stdio import stdio_server

    logger.info("启动 NotebookLM Filesystem Indexer MCP Server...")

    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
