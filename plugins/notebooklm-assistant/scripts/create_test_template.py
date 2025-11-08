#!/usr/bin/env python3
"""
创建测试用的Word文档模板

用法:
    python create_test_template.py <output.docx>
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_test_template(output_path):
    """创建一个简单的测试模板文档"""

    doc = Document()

    # 设置页面
    section = doc.sections[0]
    section.page_width = Pt(595)  # A4宽度
    section.page_height = Pt(842)  # A4高度
    section.top_margin = Pt(72)  # 1英寸
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)  # 1.25英寸
    section.right_margin = Pt(90)

    # 标题
    title = doc.add_paragraph('[项目名称]智能建造实施方案')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = '黑体'
    title.runs[0].font.size = Pt(22)
    title.runs[0].font.bold = True

    # 空行
    doc.add_paragraph()

    # 项目信息
    info = doc.add_paragraph('项目名称: [项目名称]')
    info.runs[0].font.name = '仿宋_GB2312'
    info.runs[0].font.size = Pt(14)

    info2 = doc.add_paragraph('建设地点: [建设地点]')
    info2.runs[0].font.name = '仿宋_GB2312'
    info2.runs[0].font.size = Pt(14)

    info3 = doc.add_paragraph('建设单位: [建设单位]')
    info3.runs[0].font.name = '仿宋_GB2312'
    info3.runs[0].font.size = Pt(14)

    # 空行
    doc.add_paragraph()

    # 第一章
    h1 = doc.add_paragraph('一、项目概述')
    h1.style = 'Heading 1'

    # 1.1节
    h2_1 = doc.add_paragraph('1.1 项目背景')
    h2_1.style = 'Heading 2'

    p1 = doc.add_paragraph(
        '本项目位于[建设地点],由[建设单位]投资建设。项目旨在通过智能建造技术,提升工程质量和管理水平。'
    )
    p1.runs[0].font.name = '仿宋_GB2312'
    p1.runs[0].font.size = Pt(14)
    p1.paragraph_format.first_line_indent = Pt(28)  # 首行缩进2字符
    p1.paragraph_format.line_spacing = 1.5

    # 1.2节
    h2_2 = doc.add_paragraph('1.2 项目基本信息')
    h2_2.style = 'Heading 2'

    p2 = doc.add_paragraph(
        '项目规模:_________。建设工期:_________。项目采用BIM技术、智慧工地等先进技术,打造智能建造示范工程。'
    )
    p2.runs[0].font.name = '仿宋_GB2312'
    p2.runs[0].font.size = Pt(14)
    p2.paragraph_format.first_line_indent = Pt(28)
    p2.paragraph_format.line_spacing = 1.5

    # 第二章
    h1_2 = doc.add_paragraph('二、BIM技术应用方案')
    h1_2.style = 'Heading 1'

    # 2.1节
    h2_3 = doc.add_paragraph('2.1 BIM软件选型')
    h2_3.style = 'Heading 2'

    p3 = doc.add_paragraph(
        '本项目采用Autodesk Revit作为BIM建模软件,使用Navisworks进行碰撞检测和施工模拟,通过BIM 360平台实现多方协同。'
    )
    p3.runs[0].font.name = '仿宋_GB2312'
    p3.runs[0].font.size = Pt(14)
    p3.paragraph_format.first_line_indent = Pt(28)
    p3.paragraph_format.line_spacing = 1.5

    # 2.2节
    h2_4 = doc.add_paragraph('2.2 BIM应用场景')
    h2_4.style = 'Heading 2'

    p4 = doc.add_paragraph(
        'BIM技术应用于设计优化、碰撞检测、施工模拟、进度管理等环节,实现设计施工一体化,提高工程质量和效率。'
    )
    p4.runs[0].font.name = '仿宋_GB2312'
    p4.runs[0].font.size = Pt(14)
    p4.paragraph_format.first_line_indent = Pt(28)
    p4.paragraph_format.line_spacing = 1.5

    # 第三章
    h1_3 = doc.add_paragraph('三、智慧工地管理')
    h1_3.style = 'Heading 1'

    # 3.1节
    h2_5 = doc.add_paragraph('3.1 智慧工地平台')
    h2_5.style = 'Heading 2'

    p5 = doc.add_paragraph(
        '建设智慧工地管理平台,集成人员管理、视频监控、环境监测、塔吊监控等功能,实现工地管理的数字化和智能化。'
    )
    p5.runs[0].font.name = '仿宋_GB2312'
    p5.runs[0].font.size = Pt(14)
    p5.paragraph_format.first_line_indent = Pt(28)
    p5.paragraph_format.line_spacing = 1.5

    # 保存文档
    doc.save(str(output_path))
    print(f"✅ 测试模板已创建: {output_path}")
    print(f"📏 文件大小: {output_path.stat().st_size / 1024:.1f} KB")


def main():
    if len(sys.argv) < 2:
        output_path = Path('test_template.docx')
    else:
        output_path = Path(sys.argv[1])

    create_test_template(output_path)

    print("\n💡 使用方法:")
    print(f"   python clone_format.py {output_path}")


if __name__ == '__main__':
    main()
