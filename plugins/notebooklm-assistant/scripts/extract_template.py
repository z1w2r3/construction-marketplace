#!/usr/bin/env python3
"""
从参考 Word 文档中提取完整的格式模板

功能:
1. 解析文档结构(章节、段落、表格)
2. 提取样式信息(字体、颜色、对齐方式等)
3. 提取页面设置(边距、纸张大小等)
4. 生成可重用的模板文件

用法:
    python extract_template.py <input_docx> <output_template_json>
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def extract_run_format(run):
    """提取运行(run)的格式信息"""
    format_info = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': float(run.font.size.pt) if run.font.size else None,
    }

    # 提取字体颜色
    if run.font.color and run.font.color.rgb:
        format_info['font_color'] = str(run.font.color.rgb)

    return format_info


def extract_paragraph_format(paragraph):
    """提取段落的格式信息"""
    format_info = {
        'alignment': str(paragraph.alignment) if paragraph.alignment else None,
        'line_spacing': paragraph.paragraph_format.line_spacing,
        'space_before': float(paragraph.paragraph_format.space_before.pt) if paragraph.paragraph_format.space_before else 0,
        'space_after': float(paragraph.paragraph_format.space_after.pt) if paragraph.paragraph_format.space_after else 0,
        'left_indent': float(paragraph.paragraph_format.left_indent.pt) if paragraph.paragraph_format.left_indent else 0,
        'right_indent': float(paragraph.paragraph_format.right_indent.pt) if paragraph.paragraph_format.right_indent else 0,
        'first_line_indent': float(paragraph.paragraph_format.first_line_indent.pt) if paragraph.paragraph_format.first_line_indent else 0,
    }

    # 提取段落样式
    if paragraph.style:
        format_info['style_name'] = paragraph.style.name

    return format_info


def extract_document_structure(doc_path):
    """提取文档的完整结构和格式"""
    print(f"📄 正在分析文档: {doc_path}")

    doc = Document(doc_path)

    template = {
        'metadata': {
            'source_document': str(doc_path),
            'version': '1.0'
        },
        'page_settings': {},
        'styles': {},
        'sections': [],
        'document_structure': []
    }

    # 提取页面设置
    for section in doc.sections:
        template['page_settings'] = {
            'page_width': float(section.page_width.pt),
            'page_height': float(section.page_height.pt),
            'top_margin': float(section.top_margin.pt),
            'bottom_margin': float(section.bottom_margin.pt),
            'left_margin': float(section.left_margin.pt),
            'right_margin': float(section.right_margin.pt),
        }
        break  # 只取第一个节的设置

    # 提取样式信息
    print("📐 提取样式信息...")
    for style in doc.styles:
        if style.type == WD_STYLE_TYPE.PARAGRAPH:
            try:
                style_info = {
                    'type': 'paragraph',
                    'name': style.name,
                    'based_on': style.base_style.name if style.base_style else None,
                }

                # 提取段落格式
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    pf = style.paragraph_format
                    style_info['paragraph_format'] = {
                        'alignment': str(pf.alignment) if pf.alignment else None,
                        'line_spacing': pf.line_spacing,
                        'space_before': float(pf.space_before.pt) if pf.space_before else 0,
                        'space_after': float(pf.space_after.pt) if pf.space_after else 0,
                        'left_indent': float(pf.left_indent.pt) if pf.left_indent else 0,
                        'first_line_indent': float(pf.first_line_indent.pt) if pf.first_line_indent else 0,
                    }

                # 提取字体格式
                if hasattr(style, 'font') and style.font:
                    font = style.font
                    style_info['font'] = {
                        'name': font.name,
                        'size': float(font.size.pt) if font.size else None,
                        'bold': font.bold,
                        'italic': font.italic,
                        'underline': font.underline,
                    }
                    if font.color and font.color.rgb:
                        style_info['font']['color'] = str(font.color.rgb)

                template['styles'][style.name] = style_info
            except Exception as e:
                print(f"⚠️  样式 {style.name} 提取失败: {e}")
                continue

    # 提取文档结构
    print("📋 分析文档结构...")
    current_section = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()

        if not text:
            continue

        # 判断是否为标题
        is_heading = False
        heading_level = 0

        if para.style and 'Heading' in para.style.name:
            is_heading = True
            # 提取标题级别
            try:
                heading_level = int(para.style.name.replace('Heading ', '').replace('标题 ', ''))
            except:
                heading_level = 1
        elif para.style and '标题' in para.style.name:
            is_heading = True
            try:
                heading_level = int(para.style.name.replace('标题 ', '').replace('标题', '') or '1')
            except:
                heading_level = 1

        # 提取占位符字段
        fields = extract_fields(text)

        element = {
            'index': i,
            'type': 'heading' if is_heading else 'paragraph',
            'level': heading_level if is_heading else 0,
            'text': text,
            'style': para.style.name if para.style else 'Normal',
            'format': extract_paragraph_format(para),
            'fields': fields,
        }

        # 提取运行格式(第一个运行作为参考)
        if para.runs:
            element['run_format'] = extract_run_format(para.runs[0])

        template['document_structure'].append(element)

        # 构建章节树
        if is_heading:
            if heading_level == 1:
                current_section = {
                    'title': text,
                    'level': 1,
                    'style': para.style.name if para.style else 'Heading 1',
                    'format': extract_paragraph_format(para),
                    'subsections': [],
                    'paragraphs': [],
                    'fields': fields
                }
                template['sections'].append(current_section)
            elif heading_level == 2 and current_section:
                subsection = {
                    'title': text,
                    'level': 2,
                    'style': para.style.name if para.style else 'Heading 2',
                    'format': extract_paragraph_format(para),
                    'paragraphs': [],
                    'fields': fields
                }
                current_section['subsections'].append(subsection)
        else:
            # 添加到当前章节
            if current_section:
                if current_section['subsections']:
                    current_section['subsections'][-1]['paragraphs'].append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'format': extract_paragraph_format(para),
                        'fields': fields
                    })
                else:
                    current_section['paragraphs'].append({
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                        'format': extract_paragraph_format(para),
                        'fields': fields
                    })

    # 统计信息
    stats = {
        'total_paragraphs': len(doc.paragraphs),
        'total_sections': len(template['sections']),
        'heading_levels': {},
        'total_styles': len(template['styles']),
        'total_fields': sum(len(elem.get('fields', [])) for elem in template['document_structure'])
    }

    for elem in template['document_structure']:
        if elem['type'] == 'heading':
            level = elem['level']
            stats['heading_levels'][f'level_{level}'] = stats['heading_levels'].get(f'level_{level}', 0) + 1

    template['statistics'] = stats

    return template


def extract_fields(text):
    """从文本中提取占位符字段"""
    import re
    fields = []

    # 识别各种占位符模式
    patterns = [
        (r'\[(.+?)\]', 'bracket'),           # [项目名称]
        (r'___+', 'underline'),               # _______
        (r'（\s*）', 'chinese_paren'),        # （  ）
        (r'\(\s*\)', 'paren'),                # (  )
        (r'【(.+?)】', 'double_bracket'),     # 【项目名称】
        (r'\{\{(.+?)\}\}', 'double_brace'),   # {{项目名称}}
    ]

    for pattern, field_type in patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            field_info = {
                'type': field_type,
                'position': match.start(),
                'placeholder': match.group(0),
            }

            # 提取字段名称(如果有)
            if len(match.groups()) > 0 and match.group(1):
                field_info['name'] = match.group(1).strip()
            else:
                field_info['name'] = '填空字段'

            fields.append(field_info)

    return fields


def print_summary(template):
    """打印模板摘要信息"""
    stats = template['statistics']

    print("\n" + "="*60)
    print("📊 模板提取完成")
    print("="*60)

    print(f"\n📄 源文档: {template['metadata']['source_document']}")

    print(f"\n📐 页面设置:")
    ps = template['page_settings']
    print(f"  • 纸张大小: {ps['page_width']:.1f} x {ps['page_height']:.1f} pt")
    print(f"  • 页边距: 上{ps['top_margin']:.1f} 下{ps['bottom_margin']:.1f} 左{ps['left_margin']:.1f} 右{ps['right_margin']:.1f} pt")

    print(f"\n📊 文档统计:")
    print(f"  • 总段落数: {stats['total_paragraphs']}")
    print(f"  • 章节数: {stats['total_sections']}")
    print(f"  • 样式数: {stats['total_styles']}")
    print(f"  • 识别字段: {stats['total_fields']}")

    if stats['heading_levels']:
        print(f"\n📑 标题层级:")
        for level, count in sorted(stats['heading_levels'].items()):
            print(f"  • {level}: {count} 个")

    print(f"\n📚 主要章节:")
    for i, section in enumerate(template['sections'][:5], 1):
        print(f"  {i}. {section['title']}")
        if section['fields']:
            print(f"     📝 字段: {', '.join([f['name'] for f in section['fields']])}")


def main():
    if len(sys.argv) < 2:
        print("使用方法: python extract_template.py <input.docx> [output.json]")
        print("\n示例:")
        print("  python extract_template.py template.docx")
        print("  python extract_template.py template.docx my_template.json")
        sys.exit(1)

    input_file = Path(sys.argv[1])

    if not input_file.exists():
        print(f"❌ 错误: 文件不存在 - {input_file}")
        sys.exit(1)

    # 默认输出文件名
    if len(sys.argv) >= 3:
        output_file = Path(sys.argv[2])
    else:
        output_file = input_file.stem + '_template.json'

    try:
        # 提取模板
        template = extract_document_structure(input_file)

        # 保存模板
        output_path = Path(output_file)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(template, f, ensure_ascii=False, indent=2)

        print(f"\n✅ 模板已保存: {output_path}")

        # 打印摘要
        print_summary(template)

        print("\n" + "="*60)
        print("✅ 提取完成!")
        print("="*60)

    except Exception as e:
        print(f"\n❌ 提取失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
