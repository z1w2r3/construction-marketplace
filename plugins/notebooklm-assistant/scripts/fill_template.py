#!/usr/bin/env python3
"""
使用模板和新内容生成格式完全一致的 Word 文档

功能:
1. 读取模板JSON文件
2. 读取新内容JSON文件
3. 应用模板格式生成新的Word文档

用法:
    python fill_template.py <template.json> <content.json> <output.docx>
"""

import sys
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_format(run, format_info):
    """应用运行格式"""
    if not format_info:
        return

    if 'bold' in format_info and format_info['bold'] is not None:
        run.bold = format_info['bold']

    if 'italic' in format_info and format_info['italic'] is not None:
        run.italic = format_info['italic']

    if 'underline' in format_info and format_info['underline'] is not None:
        run.underline = format_info['underline']

    if 'font_name' in format_info and format_info['font_name']:
        run.font.name = format_info['font_name']
        # 设置中文字体
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), format_info['font_name'])

    if 'font_size' in format_info and format_info['font_size']:
        run.font.size = Pt(format_info['font_size'])

    if 'font_color' in format_info and format_info['font_color']:
        try:
            # 移除 '0x' 前缀并转换为 RGB
            color_str = format_info['font_color'].replace('0x', '')
            if len(color_str) == 6:
                r = int(color_str[0:2], 16)
                g = int(color_str[2:4], 16)
                b = int(color_str[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception as e:
            print(f"⚠️  颜色设置失败: {e}")


def set_paragraph_format(paragraph, format_info):
    """应用段落格式"""
    if not format_info:
        return

    pf = paragraph.paragraph_format

    # 对齐方式
    if 'alignment' in format_info and format_info['alignment']:
        alignment_map = {
            'WD_ALIGN_PARAGRAPH.LEFT': WD_ALIGN_PARAGRAPH.LEFT,
            'WD_ALIGN_PARAGRAPH.CENTER': WD_ALIGN_PARAGRAPH.CENTER,
            'WD_ALIGN_PARAGRAPH.RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
            'WD_ALIGN_PARAGRAPH.JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
            '1': WD_ALIGN_PARAGRAPH.CENTER,
            '2': WD_ALIGN_PARAGRAPH.RIGHT,
            '3': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        alignment_str = str(format_info['alignment'])
        if alignment_str in alignment_map:
            pf.alignment = alignment_map[alignment_str]

    # 行间距
    if 'line_spacing' in format_info and format_info['line_spacing']:
        pf.line_spacing = format_info['line_spacing']

    # 段前段后间距
    if 'space_before' in format_info and format_info['space_before']:
        pf.space_before = Pt(format_info['space_before'])

    if 'space_after' in format_info and format_info['space_after']:
        pf.space_after = Pt(format_info['space_after'])

    # 缩进
    if 'left_indent' in format_info and format_info['left_indent']:
        pf.left_indent = Pt(format_info['left_indent'])

    if 'right_indent' in format_info and format_info['right_indent']:
        pf.right_indent = Pt(format_info['right_indent'])

    if 'first_line_indent' in format_info and format_info['first_line_indent']:
        pf.first_line_indent = Pt(format_info['first_line_indent'])


def apply_page_settings(doc, page_settings):
    """应用页面设置"""
    if not page_settings:
        return

    for section in doc.sections:
        if 'page_width' in page_settings:
            section.page_width = Pt(page_settings['page_width'])

        if 'page_height' in page_settings:
            section.page_height = Pt(page_settings['page_height'])

        if 'top_margin' in page_settings:
            section.top_margin = Pt(page_settings['top_margin'])

        if 'bottom_margin' in page_settings:
            section.bottom_margin = Pt(page_settings['bottom_margin'])

        if 'left_margin' in page_settings:
            section.left_margin = Pt(page_settings['left_margin'])

        if 'right_margin' in page_settings:
            section.right_margin = Pt(page_settings['right_margin'])


def create_document_from_template(template, content):
    """根据模板和内容创建新文档"""
    print("📝 开始生成文档...")

    doc = Document()

    # 应用页面设置
    if 'page_settings' in template:
        apply_page_settings(doc, template['page_settings'])
        print("✅ 页面设置已应用")

    # 生成内容
    if 'sections' in content:
        print(f"📋 生成 {len(content['sections'])} 个章节...")

        for section_idx, section_content in enumerate(content['sections']):
            # 查找对应的模板章节(通过level匹配)
            template_section = None
            if section_idx < len(template.get('sections', [])):
                template_section = template['sections'][section_idx]

            # 添加章节标题
            if 'title' in section_content:
                para = doc.add_paragraph(section_content['title'])

                # 应用标题格式
                if template_section and 'style' in template_section:
                    try:
                        para.style = template_section['style']
                    except:
                        para.style = 'Heading 1'
                else:
                    para.style = 'Heading 1'

                # 应用段落格式
                if template_section and 'format' in template_section:
                    set_paragraph_format(para, template_section['format'])

                # 应用运行格式(如果有)
                if para.runs and template_section:
                    # 从模板结构中查找对应的run格式
                    for elem in template.get('document_structure', []):
                        if elem.get('text') == template_section.get('title'):
                            if 'run_format' in elem:
                                set_run_format(para.runs[0], elem['run_format'])
                            break

            # 添加子章节
            if 'subsections' in section_content:
                for subsection in section_content['subsections']:
                    if 'title' in subsection:
                        para = doc.add_paragraph(subsection['title'])
                        para.style = 'Heading 2'

                    # 添加子章节段落
                    if 'paragraphs' in subsection:
                        for para_content in subsection['paragraphs']:
                            if isinstance(para_content, dict) and 'text' in para_content:
                                para = doc.add_paragraph(para_content['text'])

                                # 应用样式
                                if 'style' in para_content:
                                    try:
                                        para.style = para_content['style']
                                    except:
                                        para.style = 'Normal'

                                # 应用格式
                                if 'format' in para_content:
                                    set_paragraph_format(para, para_content['format'])

                                if 'run_format' in para_content and para.runs:
                                    set_run_format(para.runs[0], para_content['run_format'])
                            elif isinstance(para_content, str):
                                para = doc.add_paragraph(para_content)
                                para.style = 'Normal'

            # 添加章节段落(如果有)
            if 'paragraphs' in section_content:
                for para_content in section_content['paragraphs']:
                    if isinstance(para_content, dict) and 'text' in para_content:
                        para = doc.add_paragraph(para_content['text'])

                        # 应用样式和格式
                        if 'style' in para_content:
                            try:
                                para.style = para_content['style']
                            except:
                                para.style = 'Normal'

                        if 'format' in para_content:
                            set_paragraph_format(para, para_content['format'])

                        if 'run_format' in para_content and para.runs:
                            set_run_format(para.runs[0], para_content['run_format'])
                    elif isinstance(para_content, str):
                        para = doc.add_paragraph(para_content)
                        para.style = 'Normal'

        print("✅ 章节内容已生成")

    return doc


def validate_content(content):
    """验证内容格式"""
    if not isinstance(content, dict):
        raise ValueError("内容必须是 JSON 对象")

    if 'sections' not in content:
        raise ValueError("内容必须包含 'sections' 字段")

    if not isinstance(content['sections'], list):
        raise ValueError("'sections' 必须是列表")

    return True


def main():
    if len(sys.argv) < 4:
        print("使用方法: python fill_template.py <template.json> <content.json> <output.docx>")
        print("\n示例:")
        print("  python fill_template.py template.json new_content.json output.docx")
        sys.exit(1)

    template_file = Path(sys.argv[1])
    content_file = Path(sys.argv[2])
    output_file = Path(sys.argv[3])

    # 检查文件存在性
    if not template_file.exists():
        print(f"❌ 错误: 模板文件不存在 - {template_file}")
        sys.exit(1)

    if not content_file.exists():
        print(f"❌ 错误: 内容文件不存在 - {content_file}")
        sys.exit(1)

    try:
        # 读取模板
        print(f"📖 读取模板: {template_file}")
        with open(template_file, 'r', encoding='utf-8') as f:
            template = json.load(f)

        # 读取内容
        print(f"📖 读取内容: {content_file}")
        with open(content_file, 'r', encoding='utf-8') as f:
            content = json.load(f)

        # 验证内容
        validate_content(content)

        # 生成文档
        doc = create_document_from_template(template, content)

        # 保存文档
        print(f"\n💾 保存文档: {output_file}")
        doc.save(str(output_file))

        print("\n" + "="*60)
        print("✅ 文档生成成功!")
        print("="*60)
        print(f"\n📄 输出文件: {output_file}")
        print(f"📊 文件大小: {output_file.stat().st_size / 1024:.1f} KB")

    except Exception as e:
        print(f"\n❌ 生成失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
