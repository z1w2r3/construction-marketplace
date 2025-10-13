#!/usr/bin/env python3
"""
建筑施工文档处理 MCP 服务器
提供 Word、Excel、PDF 文档的解析和分析功能
"""

import sys
import logging
import os
from typing import Any
from datetime import datetime

# 配置日志 - 重要: 只能写到 stderr,不能写到 stdout
logging.basicConfig(
    level=logging.INFO,
    stream=sys.stderr,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 尝试导入 MCP SDK
try:
    from mcp.server import Server
    from mcp.types import Tool, TextContent
    from mcp.server.stdio import stdio_server
except ImportError as e:
    logger.error(f"MCP SDK 未安装: {e}")
    logger.error("请运行: pip install mcp python-docx openpyxl PyPDF2")
    sys.exit(1)

# 创建 MCP 服务器实例
server = Server("construction-doc-processor")

@server.list_tools()
async def list_tools() -> list[Tool]:
    """列出所有可用工具"""
    return [
        Tool(
            name="parse_word_document",
            description="解析 Word 文档,提取文本、表格和元数据",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Word 文档的绝对路径"
                    },
                    "extract_tables": {
                        "type": "boolean",
                        "description": "是否提取表格(默认 true)",
                        "default": True
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="parse_excel_document",
            description="解析 Excel 文档,提取工作表和单元格数据",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Excel 文档的绝对路径"
                    },
                    "sheet_name": {
                        "type": "string",
                        "description": "工作表名称(可选,默认读取所有)"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="parse_pdf_document",
            description="解析 PDF 文档,提取文本和元数据",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "PDF 文档的绝对路径"
                    }
                },
                "required": ["file_path"]
            }
        ),
        Tool(
            name="get_document_metadata",
            description="获取文档元数据(创建时间、修改时间、大小等)",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "文档的绝对路径"
                    }
                },
                "required": ["file_path"]
            }
        )
    ]

@server.call_tool()
async def call_tool(name: str, arguments: Any) -> list[TextContent]:
    """执行工具调用"""
    try:
        logger.info(f"调用工具: {name}, 参数: {arguments}")

        if name == "parse_word_document":
            return await parse_word_document(
                arguments["file_path"],
                arguments.get("extract_tables", True)
            )
        elif name == "parse_excel_document":
            return await parse_excel_document(
                arguments["file_path"],
                arguments.get("sheet_name")
            )
        elif name == "parse_pdf_document":
            return await parse_pdf_document(arguments["file_path"])
        elif name == "get_document_metadata":
            return await get_document_metadata(arguments["file_path"])
        else:
            raise ValueError(f"未知工具: {name}")
    except Exception as e:
        logger.error(f"工具执行错误: {e}", exc_info=True)
        return [TextContent(
            type="text",
            text=f"错误: {str(e)}"
        )]

async def parse_word_document(file_path: str, extract_tables: bool) -> list[TextContent]:
    """解析 Word 文档"""
    try:
        from docx import Document
    except ImportError:
        return [TextContent(
            type="text",
            text="错误: python-docx 未安装。请运行: pip install python-docx"
        )]

    try:
        doc = Document(file_path)

        # 提取段落文本
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        result = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "paragraph_count": len(paragraphs),
            "paragraphs_preview": paragraphs[:10],  # 只返回前10段
            "tables": []
        }

        # 提取表格
        if extract_tables and doc.tables:
            for i, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                result["tables"].append({
                    "table_index": i + 1,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "data_preview": table_data[:5]  # 只返回前5行
                })

        return [TextContent(
            type="text",
            text=f"""Word 文档解析完成:

文件: {result['file_name']}
段落数: {result['paragraph_count']}
表格数: {len(result['tables'])}

段落预览(前10段):
{chr(10).join(f'{i+1}. {p[:100]}...' if len(p) > 100 else f'{i+1}. {p}' for i, p in enumerate(result['paragraphs_preview']))}

{'表格信息:' if result['tables'] else ''}
{chr(10).join(f'表格{t["table_index"]}: {t["rows"]}行 x {t["cols"]}列' for t in result['tables'])}
"""
        )]
    except Exception as e:
        return [TextContent(
            type="text",
            text=f"Word 文档解析失败: {str(e)}"
        )]

async def parse_excel_document(file_path: str, sheet_name: str = None) -> list[TextContent]:
    """解析 Excel 文档"""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return [TextContent(
            type="text",
            text="错误: openpyxl 未安装。请运行: pip install openpyxl"
        )]

    try:
        wb = load_workbook(file_path, read_only=True, data_only=True)

        result = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "sheet_names": wb.sheetnames,
            "sheets": []
        }

        # 读取指定工作表或所有工作表
        sheets_to_read = [sheet_name] if sheet_name else wb.sheetnames[:3]  # 最多读取3个工作表

        for name in sheets_to_read:
            if name not in wb.sheetnames:
                continue
            ws = wb[name]
            data = []
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= 10:  # 只读取前10行
                    break
                data.append(list(row))

            result["sheets"].append({
                "name": name,
                "rows": ws.max_row,
                "cols": ws.max_column,
                "data_preview": data
            })

        output = f"""Excel 文档解析完成:

文件: {result['file_name']}
工作表数: {len(result['sheet_names'])}
工作表列表: {', '.join(result['sheet_names'])}

"""
        for sheet in result["sheets"]:
            output += f"\n工作表: {sheet['name']}\n"
            output += f"大小: {sheet['rows']}行 x {sheet['cols']}列\n"
            output += f"数据预览(前10行):\n"
            for i, row in enumerate(sheet['data_preview'], 1):
                row_str = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                output += f"  {i}. {row_str[:100]}...\n" if len(row_str) > 100 else f"  {i}. {row_str}\n"

        return [TextContent(type="text", text=output)]
    except Exception as e:
        return [TextContent(
            type="text",
            text=f"Excel 文档解析失败: {str(e)}"
        )]

async def parse_pdf_document(file_path: str) -> list[TextContent]:
    """解析 PDF 文档"""
    try:
        import PyPDF2
    except ImportError:
        return [TextContent(
            type="text",
            text="错误: PyPDF2 未安装。请运行: pip install PyPDF2"
        )]

    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            result = {
                "file_path": file_path,
                "file_name": os.path.basename(file_path),
                "page_count": len(reader.pages),
                "pages": []
            }

            # 提取前3页文本
            for i in range(min(3, len(reader.pages))):
                page = reader.pages[i]
                text = page.extract_text()
                result["pages"].append({
                    "page_number": i + 1,
                    "text_length": len(text),
                    "text_preview": text[:500]  # 前500字符
                })

            output = f"""PDF 文档解析完成:

文件: {result['file_name']}
总页数: {result['page_count']}

页面预览(前3页):
"""
            for page in result["pages"]:
                output += f"\n第{page['page_number']}页 (共{page['text_length']}字符):\n"
                output += f"{page['text_preview']}...\n"

            return [TextContent(type="text", text=output)]
    except Exception as e:
        return [TextContent(
            type="text",
            text=f"PDF 文档解析失败: {str(e)}"
        )]

async def get_document_metadata(file_path: str) -> list[TextContent]:
    """获取文档元数据"""
    try:
        if not os.path.exists(file_path):
            return [TextContent(
                type="text",
                text=f"错误: 文件不存在: {file_path}"
            )]

        stat = os.stat(file_path)

        result = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size": stat.st_size,
            "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
            "created_time": datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
            "modified_time": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
            "file_extension": os.path.splitext(file_path)[1]
        }

        output = f"""文档元数据:

文件名: {result['file_name']}
文件大小: {result['file_size_mb']} MB ({result['file_size']} 字节)
文件类型: {result['file_extension']}
创建时间: {result['created_time']}
修改时间: {result['modified_time']}
文件路径: {result['file_path']}
"""
        return [TextContent(type="text", text=output)]
    except Exception as e:
        return [TextContent(
            type="text",
            text=f"获取元数据失败: {str(e)}"
        )]

async def main():
    """启动 MCP 服务器"""
    logger.info("启动建筑施工文档处理 MCP 服务器...")
    async with stdio_server() as (read_stream, write_stream):
        await server.run(
            read_stream,
            write_stream,
            server.create_initialization_options()
        )

if __name__ == "__main__":
    import asyncio
    asyncio.run(main())
